from pathlib import Path
from datetime import date
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
WORK = ROOT / "work" / "doc_assets"
OUT.mkdir(exist_ok=True)
WORK.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "LAUNDRRY_CSE327_Software_Engineering_Documentation.docx"

BLUE = "3B6790"; NAVY = "1F3951"; ORANGE = "FF9D23"; GOLD = "EFB036"
LIGHT = "F2F4F7"; WHITE = "FFFFFF"; GRAY = "666666"; BLACK = "222222"

def rgb(hexv): return RGBColor.from_string(hexv)

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m, v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn('w:'+m))
        if node is None: node = OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); tblHeader = OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)

def set_table_borders(table, color="B8C2CC", size="6"):
    tblPr = table._tbl.tblPr; borders = tblPr.find(qn('w:tblBorders'))
    if borders is None: borders = OxmlElement('w:tblBorders'); tblPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = borders.find(qn('w:'+edge))
        if tag is None: tag=OxmlElement('w:'+edge); borders.append(tag)
        tag.set(qn('w:val'),'single'); tag.set(qn('w:sz'),size); tag.set(qn('w:color'),color)

def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
            tcPr=row.cells[i]._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
            if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
            tcW.set(qn('w:w'), str(int(width*1440))); tcW.set(qn('w:type'),'dxa')
    tblPr=table._tbl.tblPr; tblW=tblPr.find(qn('w:tblW'))
    if tblW is None: tblW=OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'),str(int(sum(widths)*1440))); tblW.set(qn('w:type'),'dxa')
    tblInd=tblPr.find(qn('w:tblInd'))
    if tblInd is None: tblInd=OxmlElement('w:tblInd'); tblPr.append(tblInd)
    tblInd.set(qn('w:w'),'120'); tblInd.set(qn('w:type'),'dxa')
    grid=table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        c=OxmlElement('w:gridCol'); c.set(qn('w:w'),str(int(w*1440))); grid.append(c)

def table(doc, headers, rows, widths, font=8.4):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(t); set_repeat_table_header(t.rows[0])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=str(h); set_cell_shading(c, NAVY); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: r.font.bold=True; r.font.color.rgb=rgb(WHITE); r.font.size=Pt(font)
    for ridx,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
            if ridx%2: set_cell_shading(cells[i], "F8FAFC")
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(2)
                for r in p.runs: r.font.size=Pt(font); r.font.color.rgb=rgb(BLACK)
    for row in t.rows:
        for c in row.cells: set_cell_margins(c)
    set_table_widths(t,widths); doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t

def setup_doc():
    d=Document(); sec=d.sections[0]
    sec.page_width=Inches(8.5); sec.page_height=Inches(11)
    sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
    sec.header_distance=Inches(.4); sec.footer_distance=Inches(.45)
    styles=d.styles
    normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal.font.color.rgb=rgb(BLACK)
    normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.10
    for name,size,color,before,after in [('Title',28,NAVY,0,10),('Subtitle',14,GRAY,0,10),('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,6),('Heading 3',12,NAVY,8,4)]:
        s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.color.rgb=rgb(color); s.font.bold=name!='Subtitle'
        s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    for lname in ['List Bullet','List Number']:
        s=styles[lname]; s.font.name='Calibri'; s.font.size=Pt(11); s.paragraph_format.left_indent=Inches(.5); s.paragraph_format.first_line_indent=Inches(-.25); s.paragraph_format.space_after=Pt(8); s.paragraph_format.line_spacing=1.167
    h=sec.header.paragraphs[0]; h.text='LAUNDRRY  |  Software Engineering Specification'; h.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    for r in h.runs: r.font.size=Pt(8); r.font.color.rgb=rgb(GRAY)
    f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=f.add_run('CSE327 • Team Beta • Pre-implementation baseline'); r.font.size=Pt(8); r.font.color.rgb=rgb(GRAY)
    return d

def title(doc,text,subtitle=None):
    p=doc.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(text)
    if subtitle:
        p=doc.add_paragraph(style='Subtitle'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(subtitle)

def h1(doc,text): doc.add_heading(text,level=1)
def h2(doc,text): doc.add_heading(text,level=2)
def h3(doc,text): doc.add_heading(text,level=3)
def para(doc,text,bold_prefix=None):
    p=doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold=True; p.add_run(text[len(bold_prefix):])
    else: p.add_run(text)
    return p
def bullets(doc,items):
    for x in items: doc.add_paragraph(x,style='List Bullet')
def nums(doc,items):
    for x in items: doc.add_paragraph(x,style='List Number')
def caption(doc,text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next=True
    r=p.add_run(text); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=rgb(GRAY)
def add_picture(doc,path,width=6.35):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_together=True
    shape=p.add_run().add_picture(str(path),width=Inches(width))
    descriptions={
        'use_case':'Use case diagram showing Customer, Rider, Laundry Partner, and Admin interactions with the LAUNDRRY marketplace.',
        'activity':'Swimlane activity diagram showing the order lifecycle across Customer, System, Rider, and Laundry Partner.',
        'sequence':'Sequence diagram showing order placement, rider pickup, laundry processing, return delivery, database updates, and notifications.',
        'class':'Class diagram showing User, RiderProfile, PartnerProfile, Service, Order, OrderItem, DeliveryTask, Review, and OrderStatusHistory.'}
    docPr=shape._inline.docPr; key=Path(path).stem; docPr.set('descr',descriptions.get(key,key)); docPr.set('title',key.replace('_',' ').title())

def font(size=24,bold=False):
    names=['C:/Windows/Fonts/calibrib.ttf' if bold else 'C:/Windows/Fonts/calibri.ttf','C:/Windows/Fonts/arialbd.ttf' if bold else 'C:/Windows/Fonts/arial.ttf']
    for n in names:
        try: return ImageFont.truetype(n,size)
        except OSError: pass
    return ImageFont.load_default()

def center_text(draw, xy, text, fnt, fill='#222222', spacing=3):
    box=draw.multiline_textbbox((0,0),text,font=fnt,align='center',spacing=spacing)
    draw.multiline_text((xy[0]-(box[2]-box[0])/2,xy[1]-(box[3]-box[1])/2),text,font=fnt,fill=fill,align='center',spacing=spacing)

def pbox(draw,rect,text,fill='#F2F4F7',outline='#3B6790',fs=22,radius=15):
    draw.rounded_rectangle(rect,radius=radius,fill=fill,outline=outline,width=3)
    center_text(draw,((rect[0]+rect[2])/2,(rect[1]+rect[3])/2),text,font(fs),spacing=4)

def line_arrow(draw,a,b,fill='#1F3951',width=3):
    draw.line([a,b],fill=fill,width=width)
    import math
    ang=math.atan2(b[1]-a[1],b[0]-a[0]); L=14
    for d in (.55,-.55):
        q=(b[0]-L*math.cos(ang+d),b[1]-L*math.sin(ang+d)); draw.line([b,q],fill=fill,width=width)

def save_img(img,name):
    p=WORK/name; img.save(p); return p

def diagram_use_case():
    img=Image.new('RGB',(1800,1200),'white'); d=ImageDraw.Draw(img); d.rectangle((330,80,1470,1120),outline='#3B6790',width=4); center_text(d,(900,115),'LAUNDRRY Marketplace',font(30,True),'#1F3951')
    actors={'Customer':(120,250),'Rider':(120,830),'Laundry Partner':(1680,250),'Admin':(1680,830)}
    for n,(x,y) in actors.items():
        d.ellipse((x-20,y-70,x+20,y-30),outline='black',width=3); d.line((x,y-30,x,y+45),fill='black',width=3); d.line((x-35,y,x+35,y),fill='black',width=3); d.line((x,y+45,x-30,y+95),fill='black',width=3); d.line((x,y+45,x+30,y+95),fill='black',width=3); center_text(d,(x,y+135),n,font(23,True))
    uses=[('Register / authenticate',610,235),('Browse partners & services',1110,235),('Place & track order',610,430),('Manage addresses / review',1110,430),('Accept pickup / return task',610,680),('Verify pickup / delivery',1110,680),('Manage services & prices',610,875),('Process laundry order',1110,875),('Approve users / monitor system',860,1040)]
    for t,x,y in uses: d.ellipse((x-190,y-55,x+190,y+55),fill='#FFF7E8',outline='#EFB036',width=3); center_text(d,(x,y),t,font(20))
    links=[((155,250),(420,235)),((155,250),(420,430)),((155,250),(920,430)),((155,830),(420,680)),((155,830),(920,680)),((1645,250),(800,875)),((1645,250),(1300,875)),((1645,830),(1050,1040)),((1645,830),(670,1040))]
    for a,b in links: d.line([a,b],fill='#888',width=2)
    return save_img(img,'use_case.png')

def diagram_activity():
    img=Image.new('RGB',(1400,1900),'white'); d=ImageDraw.Draw(img); lanes=['Customer','System','Rider','Partner']; centers=[175,525,875,1225]
    for i,n in enumerate(lanes): d.rectangle((i*350,0,(i+1)*350,1900),outline='#DDD',width=2); center_text(d,(centers[i],45),n,font(28,True),'#1F3951')
    nodes=[(525,120,'Start'),(175,255,'Select partner, services\nand schedule'),(525,390,'Validate and create order'),(875,525,'Accept pickup task'),(875,660,'Collect clothes'),(1225,795,'Receive and verify items'),(1225,930,'Process laundry'),(1225,1065,'Mark ready for return'),(875,1200,'Accept return task'),(875,1335,'Collect from partner'),(175,1470,'Receive and verify order'),(525,1605,'Complete order, settle\nand enable review'),(525,1740,'End')]
    last=None
    for x,y,t in nodes:
        rect=(x-135,y-42,x+135,y+42); pbox(d,rect,t,fs=21)
        if last: line_arrow(d,(last[0],last[1]+42),(x,y-42))
        last=(x,y)
    return save_img(img,'activity.png')

def diagram_sequence():
    img=Image.new('RGB',(2000,1450),'white'); d=ImageDraw.Draw(img); actors=['Customer UI','API','Order Service','Database','Rider UI','Partner UI','Notification']; xs=[140,420,700,980,1260,1540,1820]
    for x,n in zip(xs,actors): pbox(d,(x-105,30,x+105,105),n,fill='#FFF7E8',outline='#EFB036',fs=20); d.line((x,105,x,1410),fill='#AAA',width=2)
    events=[(0,1,'POST /orders'),(1,2,'validate + calculate'),(2,3,'transaction: order + task'),(3,2,'created'),(2,6,'pickup request'),(6,4,'notify'),(4,1,'accept task'),(1,2,'assign rider'),(2,3,'atomic update'),(4,1,'pickup verified'),(1,2,'status: PICKED_UP'),(4,5,'deliver clothes'),(5,1,'confirm receipt'),(1,2,'status: AT_PARTNER'),(5,1,'status: PROCESSING'),(5,1,'status: READY'),(1,2,'create return task'),(2,6,'return request'),(4,1,'accept + collect'),(4,0,'deliver + OTP'),(0,1,'confirm delivery'),(1,2,'complete + settle'),(2,3,'commit history/payment')]
    y=145
    for a,b,label in events:
        line_arrow(d,(xs[a],y),(xs[b],y)); center_text(d,((xs[a]+xs[b])/2,y-16),label,font(16),'#555'); y+=53
    return save_img(img,'sequence.png')

def diagram_class():
    img=Image.new('RGB',(2000,1400),'white'); d=ImageDraw.Draw(img)
    classes={'User':(50,100,370,370,['id','name','email','phone','role','status']),'RiderProfile':(50,640,370,880,['userId','vehicleType','availability','rating']),'PartnerProfile':(500,100,850,370,['userId','businessName','address','approvalStatus','rating']),'Service':(500,640,850,900,['id','partnerId','name','unitPrice','isActive']),'Order':(1030,100,1380,410,['id','customerId','partnerId','status','totalAmount','placedAt']),'OrderItem':(1030,640,1380,900,['id','orderId','serviceId','quantity','unitPrice']),'DeliveryTask':(1580,100,1950,430,['id','orderId','riderId','taskType','status','verificationCode']),'Review':(1580,640,1950,900,['id','orderId','rating','comment','createdAt']),'OrderStatusHistory':(1020,1060,1410,1310,['id','orderId','oldStatus','newStatus','changedAt'])}
    centers={}
    for n,(x1,y1,x2,y2,attrs) in classes.items():
        d.rectangle((x1,y1,x2,y2),fill='white',outline='#3B6790',width=3); d.rectangle((x1,y1,x2,y1+55),fill='#1F3951'); center_text(d,((x1+x2)/2,y1+27),n,font(22,True),'white')
        for i,a in enumerate(attrs): d.text((x1+18,y1+72+i*31),a,font=font(19),fill='#222')
        centers[n]=((x1+x2)//2,(y1+y2)//2)
    rel=[('User','RiderProfile'),('User','PartnerProfile'),('PartnerProfile','Service'),('User','Order'),('PartnerProfile','Order'),('Order','OrderItem'),('Service','OrderItem'),('Order','DeliveryTask'),('RiderProfile','DeliveryTask'),('Order','Review'),('Order','OrderStatusHistory')]
    for a,b in rel: d.line([centers[a],centers[b]],fill='#888',width=2)
    return save_img(img,'class.png')

doc=setup_doc()

# Cover
doc.add_paragraph().paragraph_format.space_after=Pt(55)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('LAUND'); r.font.size=Pt(34); r.bold=True; r.font.color.rgb=rgb(BLUE)
r=p.add_run('R'); r.font.size=Pt(34); r.bold=True; r.font.color.rgb=rgb(ORANGE)
r=p.add_run('RY'); r.font.size=Pt(34); r.bold=True; r.font.color.rgb=rgb(BLUE)
title(doc,'Laundry Marketplace Management System','Software Engineering Documentation and Software Requirements Specification')
doc.add_paragraph().paragraph_format.space_after=Pt(45)
table(doc,['Document field','Value'],[
 ['Course','CSE327 — Software Engineering'],['Document type','Pre-implementation SRS, UML, database design, and testing plan'],
 ['Prepared for','University project review and implementation baseline'],['Project team','Team Beta'],['Version','1.0'],['Date',date.today().isoformat()],['Status','Approved architecture — design baseline']], [2.0,4.5],9)
doc.add_page_break()

h1(doc,'Document Control')
table(doc,['Version','Date','Description','Status'],[['1.0',date.today().isoformat(),'Initial complete pre-implementation specification','Baseline']], [1,1.2,3.3,1],9)
h2(doc,'Purpose of This Document')
para(doc,'This document defines the requirements, system behavior, UML design, simplified relational database, and verification strategy for transforming the existing static LAUNDRRY website into a multi-role laundry marketplace. It is the agreed reference for design, implementation, testing, demonstration, and change control.')
h2(doc,'Intended Audience')
bullets(doc,['CSE327 course faculty and evaluators','Team Beta developers and testers','Future maintainers of the university project','Stakeholders acting as Customer, Rider, Laundry Partner, or Admin'])
h2(doc,'Definitions and Abbreviations')
table(doc,['Term','Definition'],[
 ['API','Application Programming Interface'],['SRS','Software Requirements Specification'],['RBAC','Role-Based Access Control'],['OTP','One-Time Password used to verify handover'],['Partner','A registered and approved laundry business'],['Delivery task','One logistics leg: Customer-to-Partner or Partner-to-Customer'],['MVP','Minimum Viable Product']], [1.3,5.2],9)
h2(doc,'Document Conventions')
para(doc,'Requirements use stable identifiers: FR for functional requirements, NFR for non-functional requirements, BR for business rules, UC for use cases, and TC for test cases. “Shall” indicates a mandatory requirement. Priority values are Must, Should, or Could.')

h1(doc,'1. Software Requirements Specification')
h2(doc,'1.1 Project Overview')
para(doc,'LAUNDRRY is a web-based marketplace that connects customers with verified laundry partners and riders. Customers select a partner and services, schedule pickup, and track the order. A rider transports clothes to the partner; the partner verifies and processes them; a rider returns the completed laundry. Admin supervises users, approvals, orders, and platform records.')
para(doc,'The existing HTML/CSS/JavaScript frontend provides the visual foundation: branded navigation, landing content, service cards, process steps, reviews, FAQ, and authentication screens. The implementation will preserve its orange, yellow, blue, gray, typography, illustrations, and friendly consumer-service character while replacing static data with role-aware React views and backend services.')
h2(doc,'1.2 Problem Statement')
para(doc,'The current website represents a single laundry service and cannot coordinate independent partners, riders, customers, or administrators. Booking, authentication, pricing, assignment, tracking, processing, reviews, and persistence are absent. Customers lack transparent partner choices and lifecycle visibility; laundry businesses lack a shared digital ordering channel; riders lack controlled task assignment; and administrators lack governance and auditability.')
h2(doc,'1.3 Objectives')
bullets(doc,[
 'Convert the static site into a secure multi-user laundry marketplace without discarding its established branding.',
 'Allow customers to discover partners, compare services, place orders, and follow status changes.',
 'Model pickup and return as two controlled delivery tasks that riders can accept and complete.',
 'Allow approved partners to manage services, prices, order verification, processing, and readiness.',
 'Provide administrators with approval, monitoring, moderation, and reporting functions.',
 'Apply modular design, data integrity, authorization, traceability, and systematic testing suitable for CSE327.' ])
h2(doc,'1.4 Scope')
h3(doc,'In Scope')
bullets(doc,['Responsive web application','Customer, Rider, Partner, and Admin roles','Authentication and RBAC','Partner/rider application and admin approval','Partner catalog and service pricing','One-partner-per-order checkout','Two-leg rider delivery workflow','Order status history and notifications','Cash on delivery and simulated/sandbox digital payment','Reviews, complaints, and basic administrative reports'])
h3(doc,'Out of Scope for the MVP')
bullets(doc,['Native Android/iOS applications','Machine-learning rider dispatch','Production-scale GPS route optimization','Multi-partner shopping cart','In-app voice/video communication','Cross-city logistics','Automated tax and legal compliance','Complex subscription plans and enterprise settlement'])
h2(doc,'1.5 Product Perspective and Architecture')
para(doc,'The system is a modular monolith with a React client, REST API, Node.js/Express application, and PostgreSQL database. This design separates presentation, HTTP handling, business services, authorization, and persistence while remaining practical for a university team. Optional polling or Socket.IO may provide timely status updates; the database remains the source of truth.')

h2(doc,'1.6 User Roles')
table(doc,['Role','Purpose','Principal permissions'],[
 ['Customer','Purchases laundry services','Register, manage addresses, browse approved partners, create/cancel eligible orders, track, pay, review, complain'],
 ['Rider','Performs pickup and return logistics','Apply, set availability, view/accept eligible tasks, verify handovers, update assigned task status, view earnings'],
 ['Laundry Partner','Processes customer laundry','Apply, manage business profile/services/prices, receive orders, verify items, update processing status, view settlement'],
 ['Admin','Governs the platform','Approve/suspend accounts, manage categories, monitor orders/payments/complaints, view audit and reports']], [1.15,1.65,3.7],8.5)

h2(doc,'1.7 Functional Requirements')
fr=[
('FR-01','Account registration','The system shall allow a visitor to register as Customer or submit a Rider or Partner application. Public Admin registration shall not exist.','Must'),
('FR-02','Authentication','The system shall authenticate active users and establish a secure role-aware session.','Must'),
('FR-03','Account approval','The system shall prevent pending or rejected Rider/Partner accounts from performing operational actions.','Must'),
('FR-04','Profile management','Users shall view and update permitted profile fields; customers shall manage multiple addresses.','Must'),
('FR-05','Partner discovery','Customers shall browse and filter approved, open partners serving their area.','Must'),
('FR-06','Service catalog','Partners shall create, update, enable, and disable their services, unit types, prices, and estimated processing time.','Must'),
('FR-07','Order quotation','The backend shall calculate subtotal, delivery fees, discount, and total from persisted prices before order creation.','Must'),
('FR-08','Order placement','A Customer shall create an order for exactly one approved Partner using selected services, quantities, addresses, and pickup slot.','Must'),
('FR-09','Pickup task creation','Successful order placement shall create an available Customer-to-Partner delivery task.','Must'),
('FR-10','Rider task acceptance','An available approved Rider shall accept one available task; the operation shall prevent double assignment.','Must'),
('FR-11','Pickup verification','The Rider and Customer shall confirm clothes handover using an OTP or equivalent verification.','Must'),
('FR-12','Partner receipt','The Partner shall confirm receipt and may verify quantities or propose justified adjustments.','Must'),
('FR-13','Laundry processing','The Partner shall update an accepted order from confirmed to processing and then ready for return.','Must'),
('FR-14','Return task creation','Marking an order ready shall create a Partner-to-Customer delivery task.','Must'),
('FR-15','Return delivery','An assigned Rider shall collect from the Partner and deliver to the Customer with handover verification.','Must'),
('FR-16','Status history','Every accepted order transition shall record previous status, new status, actor, time, and optional note.','Must'),
('FR-17','Order tracking','Authorized participants shall view current state and history relevant to their order or assigned task.','Must'),
('FR-18','Cancellation','A Customer shall cancel only when the configured lifecycle rules permit; reason and actor shall be recorded.','Must'),
('FR-19','Payment recording','The system shall record payment method, amount, reference, and state; client claims shall not independently confirm online payment.','Should'),
('FR-20','Notifications','The system shall create in-app notifications for material assignment and order-status events.','Should'),
('FR-21','Reviews','A Customer shall submit at most one review for a completed order.','Should'),
('FR-22','Complaints','A Customer shall open a complaint linked to an order; Admin shall record resolution.','Should'),
('FR-23','Administration','Admin shall search and manage users, approvals, orders, categories, complaints, and platform status.','Must'),
('FR-24','Reporting','Admin shall view basic totals for users, orders, statuses, revenue records, and service performance.','Should'),
('FR-25','Audit','The system shall audit approval, suspension, sensitive status, refund, and complaint-resolution actions.','Should')]
table(doc,['ID','Name','Requirement','Priority'],fr,[.65,1.25,4.05,.55],7.7)

h2(doc,'1.8 Business Rules')
br=[('BR-01','One order belongs to one Customer and one Laundry Partner.'),('BR-02','Only approved, active Partners may receive orders; only approved, available Riders may accept tasks.'),('BR-03','One order creates two delivery tasks during normal completion: outbound pickup and return.'),('BR-04','Only the backend calculates authoritative prices and totals.'),('BR-05','Only one Rider may hold an active assignment for a delivery task.'),('BR-06','Partners may change only orders assigned to their own business and only through valid transitions.'),('BR-07','A completed order may receive at most one Customer review.'),('BR-08','Admin accounts are provisioned through protected administration or database seed, never public registration.'),('BR-09','Historical order items retain price/name snapshots.'),('BR-10','All cancellations and rejected adjustments require a reason.')]
table(doc,['ID','Rule'],br,[.75,5.75],8.5)

h2(doc,'1.9 Non-functional Requirements')
nfr=[
('NFR-01','Security','Passwords shall be salted and hashed; protected APIs shall enforce authentication, role, ownership, and validation.'),
('NFR-02','Performance','Under normal university-demo load, 95% of non-report API requests should complete within 2 seconds excluding external gateway delay.'),
('NFR-03','Availability','The deployed MVP should be available during scheduled demonstrations and testing, with recoverable restart procedures.'),
('NFR-04','Reliability','Order creation, rider acceptance, status transitions, and payment updates shall use database transactions where partial updates could corrupt state.'),
('NFR-05','Usability','Core tasks shall be accessible through clear role dashboards and responsive interfaces consistent with existing LAUNDRRY branding.'),
('NFR-06','Accessibility','Forms shall have labels, keyboard navigation, visible focus, meaningful errors, and sufficient color contrast; images shall have suitable alternatives.'),
('NFR-07','Maintainability','Frontend and backend shall use modular folders, consistent naming, linting, documentation, and separation of controllers, services, and persistence.'),
('NFR-08','Scalability','Lists shall use pagination and indexed queries; application modules shall permit later extraction without requiring microservices in the MVP.'),
('NFR-09','Data integrity','Foreign keys, unique constraints, checks, enums/reference values, and immutable order snapshots shall protect relational consistency.'),
('NFR-10','Privacy','Sensitive documents and personal information shall be accessible only to authorized users and excluded from logs and client responses where unnecessary.'),
('NFR-11','Auditability','Security-sensitive and lifecycle-changing actions shall include actor and timestamp records.'),
('NFR-12','Compatibility','The responsive web client should support current desktop/mobile versions of Chrome, Edge, and Firefox.')]
table(doc,['ID','Quality','Measurable requirement'],nfr,[.75,1.0,4.75],8)

h2(doc,'1.10 User Stories and Acceptance Criteria')
stories=[
('US-C01','Customer','I want to browse partners serving my address so that I can choose a suitable laundry.','Only approved/open partners in range appear; rating, services, prices, and turnaround are visible.'),
('US-C02','Customer','I want to receive a server-calculated quotation and place an order.','Invalid services are rejected; persisted prices are used; an order number and pickup task are created atomically.'),
('US-C03','Customer','I want to track each handover and processing stage.','Current state, timestamps, and assigned rider information are visible only when applicable.'),
('US-C04','Customer','I want to review a finished service.','Review is allowed only after completion and only once per order.'),
('US-R01','Rider','I want to see eligible tasks and accept one.','Only available tasks are shown; concurrent acceptance gives the task to one rider only.'),
('US-R02','Rider','I want verified pickup and delivery steps.','OTP/equivalent must match before protected handover status is accepted.'),
('US-P01','Partner','I want to define my services and prices.','Only the owning approved Partner can change its catalog; historical orders remain unchanged.'),
('US-P02','Partner','I want to manage my processing queue.','Only valid statuses are offered; every transition is recorded.'),
('US-A01','Admin','I want to approve Rider and Partner applications.','Decision, reviewer, time, and optional reason are audited; rejected users cannot operate.'),
('US-A02','Admin','I want to monitor disputes and orders.','Admin can search records, inspect lifecycle history, and record complaint resolution.')]
table(doc,['ID','Role','Story','Acceptance criteria'],stories,[.7,.8,2.35,2.65],7.7)

h2(doc,'1.11 System Constraints')
bullets(doc,['Web-first implementation using React, REST, Node.js/Express, and PostgreSQL unless course requirements mandate equivalent technology.','Existing LAUNDRRY UI palette, logo treatment, and assets must remain recognizable.','University time, hosting budget, and team size limit production-grade dispatch, GPS, and payment functions.','External maps, SMS, email, or payment functions depend on third-party availability and credentials.','All monetary calculations use a single configured currency for the MVP.','The system serves one configured operating region initially.'])
h2(doc,'1.12 Assumptions')
bullets(doc,['Customers have a valid phone number and serviceable address.','Partners inspect received items before processing.','Riders have a suitable vehicle and can access pickup/drop-off locations.','An Admin verifies Rider and Partner documents manually.','Internet connectivity is available during operational status updates.','A Rider may handle either or both legs; the two tasks need not use the same Rider.','Cash on delivery may be used when a production payment gateway is unavailable.'])
h2(doc,'1.13 External Interface Requirements')
h3(doc,'User Interface')
para(doc,'The public interface retains existing colors and branding. Authenticated users receive role-specific navigation, dashboards, loading states, empty states, validation messages, confirmations, and responsive layouts.')
h3(doc,'Software Interfaces')
bullets(doc,['RESTful JSON API under /api/v1','PostgreSQL relational database','Optional object storage for verification documents and proof images','Optional payment, maps, SMS, or email providers through replaceable adapters'])
h3(doc,'Communication Interfaces')
para(doc,'The client communicates with the API over HTTPS. Authentication credentials use secure transport and a protected token/session strategy. Real-time updates may use polling initially and Socket.IO as an enhancement.')

h1(doc,'2. UML Design')
h2(doc,'2.1 Use Case Diagram')
para(doc,'The use case model separates customer commerce, rider logistics, partner operations, and administrative governance. Authentication is shared, while operational privileges depend on approval, role, ownership, and lifecycle state.')
caption(doc,'Figure 1 — Marketplace use case diagram'); add_picture(doc,diagram_use_case())
h3(doc,'Key Use Cases')
table(doc,['UC','Primary actor','Precondition','Success outcome'],[
 ['UC-01 Place order','Customer','Authenticated; address and approved partner available','Order, items, history, and pickup task committed'],
 ['UC-02 Accept pickup','Rider','Approved, available; task unassigned','Task atomically assigned and customer notified'],
 ['UC-03 Process laundry','Partner','Order received at owning partner','Order confirmed, processed, and marked ready'],
 ['UC-04 Complete return','Rider','Return task assigned','Verified customer handover and completed order'],
 ['UC-05 Approve application','Admin','Pending Rider/Partner application','Approved account can operate and decision is audited']], [1.25,1.05,2.0,2.2],8)

h2(doc,'2.2 Activity Diagram — Order Lifecycle')
para(doc,'The activity diagram uses swimlanes to show responsibility transfer. The system coordinates validation, state, task creation, notification, completion, and settlement; humans confirm physical handovers.')
caption(doc,'Figure 2 — End-to-end order activity'); add_picture(doc,diagram_activity(),5.6)
h3(doc,'Lifecycle Guards')
bullets(doc,['Validation failure returns the Customer to order correction without creating database records.','Failure to obtain the pickup task lock leaves the task available or assigned to the successful Rider.','A Partner adjustment pauses processing until Customer approval or authorized resolution.','Failed delivery is recorded as an exception; it does not silently complete the order.','Cancellation branches depend on current order state and configured fees.'])

h2(doc,'2.3 Sequence Diagram — Place, Pickup, Process, Return')
para(doc,'The sequence combines the four requested scenarios into one transaction-oriented interaction. The API delegates business rules to the Order Service; database commits precede notifications. UI actors never directly update the database.')
caption(doc,'Figure 3 — Main success sequence'); add_picture(doc,diagram_sequence())
h3(doc,'Sequence Responsibilities')
table(doc,['Stage','Initiator','Critical system responsibility'],[
 ['Order placement','Customer','Recalculate total, validate partner/service/slot, atomically persist order and pickup task'],
 ['Rider pickup','Rider','Atomically assign task, verify handover, append status history'],
 ['Laundry processing','Partner','Enforce ownership and transition rules, record receipt/processing/readiness'],
 ['Rider delivery','Rider/Customer','Verify final handover, complete order, record financial allocation and enable review']], [1.2,1.15,4.15],8.5)

h2(doc,'2.4 Class Diagram')
para(doc,'The conceptual class model emphasizes the simplified academic domain. User contains common identity; RiderProfile and PartnerProfile extend role-specific data. Order aggregates items, two delivery tasks, status history, and an optional review.')
caption(doc,'Figure 4 — Simplified domain class diagram'); add_picture(doc,diagram_class())
h3(doc,'Class Invariants')
bullets(doc,['Order.totalAmount is derived from immutable item price snapshots and fees.','DeliveryTask.taskType is either CUSTOMER_TO_PARTNER or PARTNER_TO_CUSTOMER.','Review requires a completed order and a matching customer.','Service belongs to exactly one Partner and can be disabled without deleting order history.','Status history is append-only in normal application operation.'])

h1(doc,'3. Simplified Database Design')
h2(doc,'3.1 Design Approach')
para(doc,'The schema intentionally focuses on Users, Laundry Partners, Riders, Services, Orders, Order Status, and Reviews. Addresses, order items, delivery tasks, and status history are included because they are necessary to express the required workflow correctly. Payments, notifications, documents, complaints, and settlements may be added as extensions without changing the core relationships.')
h2(doc,'3.2 Core Tables')
schema_rows=[
('users','id PK; name; email UQ; phone UQ; password_hash; role; account_status; created_at; updated_at','Common identity and authorization'),
('rider_profiles','user_id PK/FK; vehicle_type; vehicle_registration; availability_status; verification_status; average_rating','Rider-specific data'),
('partner_profiles','user_id PK/FK; business_name; address; latitude; longitude; service_radius_km; verification_status; is_open; average_rating','Laundry partner business data'),
('addresses','id PK; user_id FK; label; address_line; area; city; latitude; longitude; is_default','Reusable customer address'),
('services','id PK; partner_id FK; name; category; unit_type; unit_price; estimated_hours; is_active','Partner-specific catalog and prices'),
('orders','id PK; order_number UQ; customer_id FK; partner_id FK; pickup_address_id FK; return_address_id FK; status; subtotal; delivery_fee; discount; total; placed_at; completed_at','Order header and current state'),
('order_items','id PK; order_id FK; service_id FK; service_name_snapshot; quantity; unit_price_snapshot; line_total','Purchased service snapshots'),
('delivery_tasks','id PK; order_id FK; rider_id FK NULL; task_type; status; verification_code_hash; accepted_at; collected_at; delivered_at','Pickup or return logistics leg'),
('order_status_history','id PK; order_id FK; old_status; new_status; changed_by FK; note; created_at','Append-only lifecycle audit'),
('reviews','id PK; order_id UQ/FK; customer_id FK; partner_id FK; rider_id FK NULL; partner_rating; rider_rating; comment; created_at','Verified post-completion feedback')]
table(doc,['Table','Important columns','Purpose'],schema_rows,[1.25,3.65,1.6],7.6)
h2(doc,'3.3 Relationships and Cardinality')
rels=[('users → addresses','1 : many','A Customer may store several addresses.'),('users → rider_profiles','1 : 0..1','Only Rider users require a Rider profile.'),('users → partner_profiles','1 : 0..1','Only Partner users require a Partner profile.'),('partner_profiles → services','1 : many','A Partner publishes many services.'),('users(Customer) → orders','1 : many','A Customer places many orders.'),('partner_profiles → orders','1 : many','Each order is processed by one Partner.'),('orders → order_items','1 : many','Each order contains one or more items.'),('orders → delivery_tasks','1 : 2 normally','One pickup leg and one return leg.'),('rider_profiles → delivery_tasks','1 : many','A Rider may perform many tasks.'),('orders → order_status_history','1 : many','Each accepted transition is recorded.'),('orders → reviews','1 : 0..1','At most one review per completed order.')]
table(doc,['Relationship','Cardinality','Meaning'],rels,[2.25,1.0,3.25],8.1)
h2(doc,'3.4 Enumerated Values')
table(doc,['Field','Allowed MVP values'],[
 ['users.role','CUSTOMER, RIDER, PARTNER, ADMIN'],['users.account_status','PENDING, ACTIVE, SUSPENDED, REJECTED'],
 ['orders.status','PLACED, PICKUP_REQUESTED, PICKUP_ASSIGNED, PICKED_UP, AT_PARTNER, PARTNER_CONFIRMED, PROCESSING, READY_FOR_RETURN, RETURN_ASSIGNED, OUT_FOR_DELIVERY, DELIVERED, COMPLETED, CANCELLED'],
 ['delivery_tasks.task_type','CUSTOMER_TO_PARTNER, PARTNER_TO_CUSTOMER'],['delivery_tasks.status','AVAILABLE, ACCEPTED, ARRIVED, COLLECTED, DELIVERED, FAILED'],['services.unit_type','ITEM, KG']], [2.1,4.4],8)
h2(doc,'3.5 Keys, Constraints, and Indexes')
bullets(doc,['Unique indexes on normalized user email, phone, and order number.','Foreign keys prevent orphan profiles, services, orders, items, tasks, history, and reviews.','CHECK constraints require positive quantity/prices, ratings from 1 to 5, and nonnegative totals.','Unique order_id on reviews enforces one review per order.','A partial unique index or transactional lock prevents more than one active Rider assignment per task.','Indexes on orders(customer_id, created_at), orders(partner_id, status), delivery_tasks(status, rider_id), services(partner_id, is_active), and history(order_id, created_at).','Soft-disable services and suspend users rather than deleting records required by order history.'])
h2(doc,'3.6 Transaction Boundaries')
table(doc,['Operation','Records committed together'],[
 ['Place order','Order, order items, initial history, pickup delivery task'],['Accept task','Task rider assignment, task status, order state/history'],['Partner marks ready','Order state/history, return delivery task'],['Confirm final delivery','Task state, order state/history, payment/earning records when implemented']], [1.75,4.75],8.5)

h1(doc,'4. Testing Plan')
h2(doc,'4.1 Test Strategy')
para(doc,'Testing follows a risk-based pyramid. Pure business rules receive fast unit tests; database and API boundaries receive integration tests; user-visible behavior receives black-box and end-to-end tests. Every Must requirement must have at least one passing verification before release. Defects are logged with severity, environment, reproduction steps, expected result, actual result, and evidence.')
h2(doc,'4.2 Test Environment')
table(doc,['Area','Proposed environment'],[
 ['Frontend','React test runner, browser automation, Chrome/Edge responsive views'],['Backend','Node test runner with isolated environment variables'],['Database','Dedicated PostgreSQL test database reset between suites'],['API','HTTP integration client against test server'],['Fixtures','Seeded Admin, Customer, approved/pending Rider, approved/pending Partner, services, addresses'],['External services','Mocks/sandbox for notification and payment adapters']], [1.4,5.1],8.5)
h2(doc,'4.3 Unit Testing')
bullets(doc,['Order quotation and total calculation','Valid and invalid order-state transitions','Cancellation eligibility and fee rule','Role/ownership policy decisions','Partner service validation','OTP comparison and expiration behavior','Rating and review eligibility','Task eligibility and earnings calculation','Response mapping that removes sensitive fields'])
h2(doc,'4.4 Integration Testing')
bullets(doc,['Registration, login, refresh/logout, and protected route access','Rider/Partner application approval persistence','Order creation transaction and rollback','Concurrent Rider acceptance of the same task','Partner ownership and lifecycle updates','Return-task generation exactly once','Status-history append behavior','Review uniqueness and completion requirement','Payment callback signature and idempotency when implemented'])
h2(doc,'4.5 Black-box Testing')
para(doc,'Black-box tests derive inputs from requirements without inspecting implementation. Equivalence partitioning covers valid/invalid roles, statuses, quantities, and ratings. Boundary-value analysis covers zero/maximum quantities, rating limits, time slots, and field lengths. Decision tables cover cancellation and authorization. State-transition tests cover the order lifecycle.')
h2(doc,'4.6 Entry and Exit Criteria')
table(doc,['Gate','Criteria'],[
 ['Test entry','Approved requirements; runnable build; migrated test database; stable test data; known environment documented'],
 ['Feature exit','All Must acceptance criteria pass; no open Critical/High defect; changed modules have unit and integration coverage'],
 ['Release exit','End-to-end success workflow passes; security-negative cases pass; backup/setup/demo instructions verified; residual risks documented']], [1.25,5.25],8.3)
h2(doc,'4.7 Detailed Test Cases')
tcs=[
('TC-01','Customer registration','Unique valid name/email/phone/password','Submit registration','ACTIVE Customer created; password not stored plainly','FR-01'),
('TC-02','Reject duplicate identity','Existing email','Register with same email','Validation error; no new user','FR-01'),
('TC-03','Block public Admin registration','Role payload ADMIN','Submit registration API request','403/validation failure; no Admin created','FR-01'),
('TC-04','Login success','Active Customer','Submit correct credentials','Authenticated session and Customer role returned','FR-02'),
('TC-05','Login failure','Existing user','Submit wrong password','Generic authentication error; no session','FR-02,NFR-01'),
('TC-06','Pending Rider blocked','Pending Rider','Request available tasks','Access denied','FR-03'),
('TC-07','Admin approves Partner','Pending Partner, Admin logged in','Approve application','Partner ACTIVE; audit entry created','FR-03,FR-25'),
('TC-08','Partner list filtering','Approved/open and rejected partners','Search service area','Only eligible partners returned','FR-05'),
('TC-09','Create service','Approved Partner','Submit valid name/unit/price','Service belongs to Partner and is visible','FR-06'),
('TC-10','Reject invalid price','Approved Partner','Submit zero/negative price','Validation error; no service persisted','FR-06,NFR-09'),
('TC-11','Server quotation','Customer selects active services','Request quote','Correct subtotal, fees, discount, total','FR-07'),
('TC-12','Reject manipulated total','Client sends lower total','Place order','Server ignores/rejects client total and calculates authoritative amount','FR-07'),
('TC-13','Place valid order','Customer, address, partner, services','Confirm checkout','Order/items/history/pickup task committed','FR-08,FR-09'),
('TC-14','Rollback invalid order','One inactive service','Place order','No partial order/task remains','FR-08,NFR-04'),
('TC-15','Concurrent task acceptance','Two approved Riders, one task','Accept simultaneously','Exactly one Rider succeeds','FR-10'),
('TC-16','Pickup OTP success','Assigned Rider, valid OTP','Confirm collection','Task COLLECTED; order PICKED_UP; history appended','FR-11,FR-16'),
('TC-17','Pickup OTP failure','Assigned Rider, wrong OTP','Confirm collection','Rejected; state unchanged','FR-11'),
('TC-18','Partner confirms receipt','Owning Partner, PICKED_UP order','Confirm receipt','Order AT_PARTNER/PARTNER_CONFIRMED per transition','FR-12'),
('TC-19','Wrong Partner access','Different Partner','Open/update order','404/403; no information or state change','FR-12,NFR-01'),
('TC-20','Processing lifecycle','Owning Partner','Confirm → Processing → Ready','Transitions accepted and timestamped','FR-13,FR-16'),
('TC-21','Invalid status jump','Partner at AT_PARTNER','Set COMPLETED','Rejected; state unchanged','FR-13'),
('TC-22','Generate return once','Order PROCESSING','Mark ready twice/retry request','Exactly one return task exists','FR-14,NFR-04'),
('TC-23','Return delivery','Assigned Rider and Customer','Collect, deliver, verify OTP','Task delivered; order completed','FR-15'),
('TC-24','Tracking authorization','Unrelated Customer','Request another order tracking','Access denied/no sensitive data','FR-17,NFR-01'),
('TC-25','Early cancellation','Customer, unassigned order','Cancel with reason','Order CANCELLED; task closed; history records actor/reason','FR-18'),
('TC-26','Late cancellation','Order PROCESSING','Customer requests cancellation','Automatic cancellation rejected; complaint/admin path offered','FR-18'),
('TC-27','Review completed order','Owner Customer, completed order','Submit ratings 1–5','Review saved; aggregates recalculated','FR-21'),
('TC-28','Duplicate review','Reviewed completed order','Submit second review','Conflict response; only one review remains','FR-21'),
('TC-29','Responsive order flow','Mobile viewport','Browse through checkout','No clipped controls; keyboard/form interaction works','NFR-05,NFR-06'),
('TC-30','API response performance','Representative seeded dataset','Execute standard list/detail calls','95% complete within 2 seconds in test profile','NFR-02')]
table(doc,['ID','Scenario','Precondition/Input','Action','Expected result','Trace'],tcs,[.55,1.1,1.25,1.05,2.05,.5],6.8)
h2(doc,'4.8 Requirement Traceability Summary')
table(doc,['Requirement area','Primary verification'],[
 ['Identity and approval','TC-01 through TC-07'],['Partner catalog','TC-08 through TC-10'],['Quotation and ordering','TC-11 through TC-14'],['Pickup logistics','TC-15 through TC-17'],['Partner processing','TC-18 through TC-22'],['Return and tracking','TC-23 through TC-24'],['Cancellation and feedback','TC-25 through TC-28'],['Quality attributes','TC-05, TC-14, TC-15, TC-19, TC-24, TC-29, TC-30']], [2.3,4.2],8.5)
h2(doc,'4.9 Defect Severity')
table(doc,['Severity','Meaning','Example'],[
 ['Critical','Security/data loss or no usable core workflow','Unauthorized Admin creation; corrupted orders'],['High','Core scenario fails without reasonable workaround','No Rider can accept pickup'],['Medium','Feature defect with workaround','Incorrect filter ordering'],['Low','Cosmetic or minor usability issue','Noncritical spacing inconsistency']], [1.0,2.55,2.95],8.3)

h1(doc,'5. Implementation Governance')
h2(doc,'5.1 Recommended Delivery Order')
nums(doc,['Finalize SRS, diagrams, ER model, and acceptance criteria.','Migrate existing public UI to reusable React components.','Implement database migrations, seed Admin, authentication, and RBAC.','Implement Partner approval, catalog, and discovery.','Implement Customer quotation and order transaction.','Implement pickup task, Rider acceptance, and verified handover.','Implement Partner verification and processing lifecycle.','Implement return task, verified delivery, and completion.','Add reviews, notifications, complaints, payment sandbox, and reporting.','Complete testing, security review, deployment, and demonstration documentation.'])
h2(doc,'5.2 Change Control')
para(doc,'A change affecting scope, roles, order states, database relationships, authorization, or acceptance criteria shall be recorded as a change request. The team assesses impact on requirements, design, implementation, tests, schedule, and demo. Approved changes increment the document version and update traceability before implementation.')
h2(doc,'5.3 Definition of Done')
bullets(doc,['Requirement and acceptance criteria are implemented.','Code follows agreed module boundaries and passes linting/review.','Unit and integration tests pass.','Relevant negative authorization and invalid-state tests pass.','Database migration and rollback implications are documented.','Responsive UI includes loading, empty, error, and success states.','API/OpenAPI and user-facing documentation are updated.','No unresolved Critical or High defect remains.'])

h1(doc,'Appendix A — Order Transition Authority')
trans=[
('PLACED','PICKUP_REQUESTED','System','Order transaction succeeds'),('PICKUP_REQUESTED','PICKUP_ASSIGNED','System/Rider service','Eligible Rider wins atomic assignment'),('PICKUP_ASSIGNED','PICKED_UP','Assigned Rider','Valid handover verification'),('PICKED_UP','AT_PARTNER','Partner/assigned Rider','Partner confirms receipt'),('AT_PARTNER','PARTNER_CONFIRMED','Owning Partner','Items verified'),('PARTNER_CONFIRMED','PROCESSING','Owning Partner','Processing started'),('PROCESSING','READY_FOR_RETURN','Owning Partner','Service completed'),('READY_FOR_RETURN','RETURN_ASSIGNED','System/Rider service','Return task assigned'),('RETURN_ASSIGNED','OUT_FOR_DELIVERY','Assigned Rider','Collected from Partner'),('OUT_FOR_DELIVERY','DELIVERED','Assigned Rider/Customer','Valid final handover'),('DELIVERED','COMPLETED','System','Completion transaction succeeds'),('Eligible early states','CANCELLED','Customer/Admin/System','Cancellation policy satisfied')]
table(doc,['From','To','Authorized actor','Guard'],trans,[1.35,1.35,1.55,2.25],8)
h1(doc,'Appendix B — Role Permission Matrix')
perms=[('Manage own profile','Yes','Yes','Yes','Yes'),('Browse marketplace','Yes','No','No','Yes'),('Place order','Yes','No','No','No'),('Accept delivery task','No','Yes','No','No'),('Update assigned delivery','No','Yes','No','Admin oversight'),('Manage partner services','No','No','Own only','Yes'),('Update processing status','No','No','Own orders','Admin exception'),('Approve accounts','No','No','No','Yes'),('Monitor all orders','Own only','Assigned only','Own only','Yes'),('Submit review','Completed own order','No','No','Moderate'),('Resolve complaint','No','No','No','Yes')]
table(doc,['Capability','Customer','Rider','Partner','Admin'],perms,[2.1,1.1,1.1,1.1,1.1],7.8)

doc.core_properties.title='LAUNDRRY Laundry Marketplace Management System — Software Engineering Documentation'
doc.core_properties.subject='CSE327 pre-implementation SRS, UML, database design, and testing plan'
doc.core_properties.author='Team Beta'
doc.core_properties.keywords='CSE327, SRS, UML, laundry marketplace, software engineering'
doc.save(DOCX)
print(DOCX)
